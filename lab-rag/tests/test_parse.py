from pathlib import Path

import pytest

from labrag.parse import ParsedDoc, UnsupportedFileType, clean_text, is_supported, parse_file


def make_pdf(path: Path, pages: list[str], title_line: str | None = None, metadata: dict | None = None):
    import pymupdf

    doc = pymupdf.open()
    for i, body in enumerate(pages):
        page = doc.new_page()
        y = 72
        if i == 0 and title_line:
            page.insert_text((72, y), title_line, fontsize=20)
            y += 40
        for line in body.split("\n"):
            page.insert_text((72, y), line, fontsize=10)
            y += 14
    if metadata:
        doc.set_metadata(metadata)
    doc.save(str(path))
    doc.close()


def test_clean_text_fixes_hyphenation_and_whitespace():
    assert clean_text("hydro-\ndynamic   swim-\r\nming\n\n\n\nnext") == "hydrodynamic swimming\n\nnext"


def test_supported_extensions():
    assert is_supported(Path("a.PDF"))
    assert is_supported(Path("notes.md"))
    assert not is_supported(Path("data.csv"))
    assert not is_supported(Path("movie.mp4"))


def test_unsupported_raises(tmp_path):
    p = tmp_path / "x.xyz"
    p.write_text("hi")
    with pytest.raises(UnsupportedFileType):
        parse_file(p)


def test_pdf_title_year_doi_from_text(tmp_path):
    body = (
        "Received 3 March 2019; accepted 9 June 2019\n"
        "doi:10.1234/abcd.5678.\n"
        + "White sharks aggregate seasonally near pinniped colonies. " * 8
    )
    p = tmp_path / "paper.pdf"
    make_pdf(p, [body, "Methods. " * 60], title_line="Seasonal aggregation of white sharks")
    doc = parse_file(p)
    assert doc.n_pages == 2
    assert doc.title == "Seasonal aggregation of white sharks"
    assert doc.year == 2019
    assert doc.doi == "10.1234/abcd.5678"
    assert not doc.needs_ocr
    assert "aggregate seasonally" in doc.text


def test_pdf_metadata_title_and_author_preferred(tmp_path):
    p = tmp_path / "x.pdf"
    make_pdf(p, ["Body text. " * 50], metadata={"title": "A Metadata Title", "author": "Jorgensen S"})
    doc = parse_file(p)
    assert doc.title == "A Metadata Title"
    assert doc.authors == "Jorgensen S"


def test_pdf_junk_metadata_ignored(tmp_path):
    p = tmp_path / "Smith_2020_Leopard shark movement.pdf"
    make_pdf(p, ["Body text. " * 50], metadata={"title": "Microsoft Word", "author": ""})
    doc = parse_file(p)
    assert doc.title != "Microsoft Word"
    assert doc.authors == "Smith"
    assert doc.year == 2020


def test_scanned_pdf_flagged_for_ocr(tmp_path):
    p = tmp_path / "scan.pdf"
    make_pdf(p, ["", ""])
    doc = parse_file(p)
    assert doc.needs_ocr
    assert doc.title == "scan"


def test_filename_author_year_with_et_al(tmp_path):
    p = tmp_path / "Jorgensen et al. 2010 - Philopatry and migration.txt"
    p.write_text("Some notes about white sharks.\n")
    doc = parse_file(p)
    assert doc.authors == "Jorgensen"
    assert doc.year == 2010


def test_markdown_title_from_heading(tmp_path):
    p = tmp_path / "notes.md"
    p.write_text("# Lab meeting 2024-05-01\n\nWe discussed tagging.\n")
    doc = parse_file(p)
    assert doc.title == "Lab meeting 2024-05-01"
    assert doc.year == 2024
    assert "tagging" in doc.text
    assert "#" not in doc.text


def test_html_strips_scripts(tmp_path):
    p = tmp_path / "page.html"
    p.write_text("<html><head><title>Shark page</title><script>var x=1;</script></head>"
                 "<body><p>Sevengill sharks</p><style>p{}</style><p>eat rays</p></body></html>")
    doc = parse_file(p)
    assert doc.title == "Shark page"
    assert "var x" not in doc.text
    assert "Sevengill sharks" in doc.text and "eat rays" in doc.text


def test_docx(tmp_path):
    import docx

    d = docx.Document()
    d.add_heading("Bat ray foraging", level=1)
    d.add_paragraph("Bat rays excavate pits in the mudflat.")
    d.core_properties.author = "Sambold D"
    p = tmp_path / "rays.docx"
    d.save(str(p))
    doc = parse_file(p)
    assert doc.title == "Bat ray foraging"
    assert doc.authors == "Sambold D"
    assert "excavate pits" in doc.text
    assert isinstance(doc, ParsedDoc)


def test_year_not_taken_from_doi(tmp_path):
    p = tmp_path / "paper.pdf"
    make_pdf(p, ["doi:10.1098/rspb.2009.1155\nPublished 2012. " + "Body text. " * 40])
    doc = parse_file(p)
    assert doc.doi == "10.1098/rspb.2009.1155"
    assert doc.year == 2012


def test_pdf_blocks_become_paragraphs_and_references_are_dropped(tmp_path):
    import pymupdf

    p = tmp_path / "refs.pdf"
    d = pymupdf.open()
    page = d.new_page()
    y = 72
    for para in ["Intro paragraph about sharks. " * 6, "Methods paragraph about tags. " * 6, "Results paragraph. " * 6]:
        page.insert_textbox(pymupdf.Rect(72, y, 540, y + 90), para, fontsize=9)
        y += 100
    page.insert_text((72, y + 10), "References", fontsize=11)
    page.insert_text((72, y + 30), "Smith J (2019) Some cited paper. Journal 1:1-10.", fontsize=9)
    d.save(str(p))
    d.close()
    doc = parse_file(p)
    assert "\n\n" in doc.pages[0]  # blocks separated by blank lines
    from labrag.chunk import chunk_pages

    chunks = chunk_pages(doc.pages)
    assert chunks and not any("Some cited paper" in c.text for c in chunks)


def test_two_column_pdf_keeps_reading_order(tmp_path):
    import pymupdf

    p = tmp_path / "twocol.pdf"
    d = pymupdf.open()
    page = d.new_page()
    left = ["LEFT-ONE first column top paragraph. " * 3, "LEFT-TWO first column second paragraph. " * 3]
    right = ["RIGHT-ONE second column top paragraph. " * 3, "RIGHT-TWO second column second paragraph. " * 3]
    y = 72
    for para in left:
        page.insert_textbox(pymupdf.Rect(60, y, 290, y + 80), para, fontsize=9)
        y += 90
    y = 72
    for para in right:
        page.insert_textbox(pymupdf.Rect(310, y, 540, y + 80), para, fontsize=9)
        y += 90
    d.save(str(p))
    d.close()
    text = parse_file(p).pages[0]
    order = [text.index(k) for k in ("LEFT-ONE", "LEFT-TWO", "RIGHT-ONE", "RIGHT-TWO")]
    assert order == sorted(order), text
