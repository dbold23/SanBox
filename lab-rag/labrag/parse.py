"""Turn a document file into plain text plus the metadata needed for a citation.

Supported: PDF (text layer), DOCX, TXT, Markdown, HTML. Everything here is
heuristic and deliberately simple: we want a title, a year, and if possible a
first author and DOI, so that answers can say "[Smith 2019]" and link back to
the file. When a heuristic fails we fall back to the filename, which in most
labs already carries the author and year.
"""

from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field
from datetime import date
from html.parser import HTMLParser
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".html", ".htm"}

# A scanned PDF with no text layer yields almost nothing; we flag it instead of
# indexing an empty document.
MIN_TEXT_CHARS = 200

_DOI_RE = re.compile(r"\b(10\.\d{4,9}/[^\s\"'<>]+)")
_YEAR_RE = re.compile(r"(?<!\d)((?:19|20)\d{2})(?!\d)")
# Author_Year_Title.pdf or "Smith et al. 2019 - Title.pdf" style filenames
_FILENAME_AUTHOR_YEAR_RE = re.compile(
    r"^(?P<author>[A-Z][A-Za-z'\-]+)(?:\s*et\s*al\.?)?[\s_\-]+(?P<year>(?:19|20)\d{2})\b"
)


@dataclass
class ParsedDoc:
    text: str
    pages: list[str] = field(default_factory=list)
    title: str | None = None
    authors: str | None = None
    year: int | None = None
    doi: str | None = None
    needs_ocr: bool = False

    @property
    def n_pages(self) -> int:
        return len(self.pages)


class UnsupportedFileType(ValueError):
    pass


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTENSIONS


def parse_file(path: Path) -> ParsedDoc:
    """Parse one file. Raises UnsupportedFileType for unknown extensions."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        doc = _parse_pdf(path)
    elif ext == ".docx":
        doc = _parse_docx(path)
    elif ext in {".html", ".htm"}:
        doc = _parse_html(path)
    elif ext in {".txt", ".md", ".markdown"}:
        doc = _parse_text(path)
    else:
        raise UnsupportedFileType(f"{path.name}: unsupported file type {ext!r}")

    _fill_metadata_from_filename_and_text(doc, path)
    return doc


# --------------------------------------------------------------------------- PDF


def _parse_pdf(path: Path) -> ParsedDoc:
    import pymupdf  # imported lazily so the rest of the package works without it

    pages: list[str] = []
    title = None
    authors = None
    year = None
    with pymupdf.open(str(path)) as pdf:
        for page in pdf:
            pages.append(_page_text(page))
        meta = pdf.metadata or {}
        title = _clean_meta(meta.get("title"))
        authors = _clean_meta(meta.get("author"))
        created = meta.get("creationDate") or ""
        m = re.search(r"D:((?:19|20)\d{2})", created)
        if m:
            year = int(m.group(1))
        if not title and len(pdf) > 0:
            title = _largest_font_line(pdf[0])

    text = "\n\n".join(p for p in pages if p)
    needs_ocr = len(text) < MIN_TEXT_CHARS and len(pages) > 0
    return ParsedDoc(text=text, pages=pages, title=title, authors=authors, year=year, needs_ocr=needs_ocr)


def _page_text(page) -> str:
    """Page text with real paragraph breaks: one blank line between layout blocks."""
    try:
        blocks = page.get_text("blocks")
    except Exception:  # pragma: no cover - pymupdf internals
        return clean_text(page.get_text("text"))
    parts = []
    # Keep pymupdf's natural block order: it follows the PDF content stream, which is the
    # reading order for two-column journal layouts. Sorting by position would interleave columns.
    for b in (b for b in blocks if len(b) > 6 and b[6] == 0):
        txt = clean_text(b[4])
        if txt:
            parts.append(" ".join(txt.split("\n")) if "\n\n" not in txt else txt)
    return "\n\n".join(parts)


def _largest_font_line(page) -> str | None:
    """Guess the title: the largest-font text in the top 60% of page 1."""
    try:
        layout = page.get_text("dict")
    except Exception:  # pragma: no cover - pymupdf internals
        return None
    page_height = page.rect.height or 1
    best_size = 0.0
    best_lines: list[str] = []
    for block in layout.get("blocks", []):
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            y = line.get("bbox", [0, 0, 0, 0])[1]
            if y > page_height * 0.6:
                continue
            size = max(s.get("size", 0) for s in spans)
            txt = clean_text("".join(s.get("text", "") for s in spans))
            if len(txt) < 4 or len(txt) > 300:
                continue
            if size > best_size + 0.5:
                best_size, best_lines = size, [txt]
            elif abs(size - best_size) <= 0.5:
                best_lines.append(txt)
    title = " ".join(best_lines).strip()
    if not title or len(title) < 8:
        return None
    return title[:300]


# -------------------------------------------------------------------------- DOCX


def _parse_docx(path: Path) -> ParsedDoc:
    import docx  # python-docx

    document = docx.Document(str(path))
    paragraphs = [clean_text(p.text) for p in document.paragraphs]
    text = "\n\n".join(p for p in paragraphs if p)
    props = document.core_properties
    title = _clean_meta(props.title)
    if not title:
        for p in document.paragraphs:
            if p.style is not None and p.style.name.lower().startswith(("heading", "title")) and p.text.strip():
                title = clean_text(p.text)
                break
    authors = _clean_meta(props.author)
    year = props.created.year if props.created else None
    return ParsedDoc(text=text, pages=[text], title=title, authors=authors, year=year)


# --------------------------------------------------------------------- HTML/TEXT


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.title_parts: list[str] = []
        self._skip = 0
        self._in_title = False

    def handle_starttag(self, tag, attrs):
        if tag in {"script", "style", "noscript"}:
            self._skip += 1
        elif tag == "title":
            self._in_title = True
        elif tag in {"p", "div", "br", "li", "h1", "h2", "h3", "h4", "tr", "section", "article"}:
            self.parts.append("\n")

    def handle_endtag(self, tag):
        if tag in {"script", "style", "noscript"} and self._skip:
            self._skip -= 1
        elif tag == "title":
            self._in_title = False
        elif tag in {"p", "div", "li", "h1", "h2", "h3", "h4", "tr", "section", "article"}:
            self.parts.append("\n")

    def handle_data(self, data):
        if self._skip:
            return
        if self._in_title:
            self.title_parts.append(data)
        self.parts.append(data)


def _parse_html(path: Path) -> ParsedDoc:
    raw = path.read_text(encoding="utf-8", errors="replace")
    extractor = _TextExtractor()
    extractor.feed(raw)
    text = clean_text("".join(extractor.parts))
    title = clean_text("".join(extractor.title_parts)) or None
    return ParsedDoc(text=text, pages=[text], title=title)


def _parse_text(path: Path) -> ParsedDoc:
    raw = path.read_text(encoding="utf-8", errors="replace")
    body = re.sub(r"^#{1,6}\s+", "", raw, flags=re.M) if path.suffix.lower() in {".md", ".markdown"} else raw
    text = clean_text(body)
    title = None
    for line in raw.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip() or None
        elif len(stripped) <= 200:
            title = stripped
        break
    return ParsedDoc(text=text, pages=[text], title=title)


# ----------------------------------------------------------------------- helpers


def clean_text(s: str) -> str:
    """Normalise whitespace, fix hyphenated line breaks, keep paragraph breaks."""
    if not s:
        return ""
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"(\w)-\n(\w)", r"\1\2", s)  # "hydro-\ndynamic" -> "hydrodynamic"
    s = re.sub(r"[ \t\f\v]+", " ", s)
    s = re.sub(r" *\n *", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def _clean_meta(value) -> str | None:
    if not value:
        return None
    value = clean_text(str(value))
    # PDF producers love to put junk here
    junk = {"untitled", "microsoft word", "title", "pdf", "document"}
    lowered = value.lower()
    if lowered in junk or lowered.endswith((".doc", ".docx", ".indd", ".dvi", ".tex")) or len(value) < 3:
        return None
    return value[:300]


def title_from_filename(path: Path) -> str:
    stem = path.stem
    stem = re.sub(r"[_]+", " ", stem)
    stem = re.sub(r"\s+", " ", stem).strip()
    return stem or path.name


def _fill_metadata_from_filename_and_text(doc: ParsedDoc, path: Path) -> None:
    head = "\n".join(doc.pages[:2]) if doc.pages else doc.text[:20000]

    if not doc.doi:
        m = _DOI_RE.search(head)
        if m:
            doc.doi = m.group(1).rstrip(".,;:)")

    fm = _FILENAME_AUTHOR_YEAR_RE.match(path.stem.replace("_", " "))
    if fm:
        if not doc.authors:
            doc.authors = fm.group("author")
        if not doc.year:
            doc.year = int(fm.group("year"))

    if not doc.year:
        doc.year = _guess_year(path.stem) or _guess_year(head[:6000])

    if not doc.title:
        doc.title = title_from_filename(path)


def _guess_year(text: str) -> int | None:
    this_year = date.today().year
    text = _DOI_RE.sub(" ", text)  # a DOI like 10.1098/rspb.2009.1155 is not a publication year
    years = [int(y) for y in _YEAR_RE.findall(text) if 1900 <= int(y) <= this_year + 1]
    if not years:
        return None
    # the most common plausible year; ties go to the most recent
    counts = Counter(years)
    return max(counts, key=lambda y: (counts[y], y))
